from io import BytesIO

from django.http import HttpResponse


DOCX_CONTENT_TYPE = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'


def exportar_documento_conceptos_docx(documento):
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    doc.add_heading('Documento de conceptos', level=1)
    datos = [
        ('Folio', documento.folio),
        ('Estado', documento.status),
        ('Fecha de creacion', documento.created_at.strftime('%Y-%m-%d %H:%M')),
        ('Usuario', documento.usuario.username if documento.usuario else ''),
    ]
    if documento.observaciones:
        datos.append(('Observaciones', documento.observaciones))
    for etiqueta, valor in datos:
        parrafo = doc.add_paragraph()
        parrafo.add_run(f'{etiqueta}: ').bold = True
        parrafo.add_run(str(valor))

    conceptos = documento.conceptos.order_by('orden', 'pk')
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    headers = [
        'Identificación',
        'Descripcion',
        'Cantidad',
        'Precio unitario',
        'Total concepto',
    ]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
        for paragraph in table.rows[0].cells[index].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for concepto in conceptos:
        cells = table.add_row().cells
        _agregar_identificacion(cells[0], concepto)
        cells[1].text = concepto.descripcion
        cells[2].text = _decimal_legible(concepto.cantidad)
        cells[3].text = _decimal_legible(concepto.precio_unitario)
        cells[4].text = _decimal_legible(concepto.total_concepto)
        for index in (2, 3, 4):
            for paragraph in cells[index].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    total = doc.add_paragraph()
    total.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    total.add_run('Total general: ').bold = True
    total.add_run(_decimal_legible(documento.total))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    response = HttpResponse(buffer.read(), content_type=DOCX_CONTENT_TYPE)
    response['Content-Disposition'] = f'attachment; filename="conceptos_{documento.folio}.docx"'
    return response


def _decimal_legible(valor):
    texto = f'{valor:.6f}'
    return texto.rstrip('0').rstrip('.') if '.' in texto else texto


def _agregar_identificacion(cell, concepto):
    cell.text = ''
    bloques = (
        ('Número de parte.', concepto.numero_parte),
        ('Modelo.', concepto.modelo),
        ('Serie.', concepto.serie),
    )
    for indice, (etiqueta, valor) in enumerate(bloques):
        if indice:
            cell.add_paragraph()
        etiqueta_parrafo = cell.add_paragraph()
        etiqueta_parrafo.add_run(etiqueta).bold = True
        cell.add_paragraph(valor or '—')
