import json
from decimal import Decimal
from io import BytesIO, StringIO

from django.contrib.auth import get_user_model
from django.contrib.auth.models import Permission
from django.core.files.uploadedfile import SimpleUploadedFile
from django.core.management import call_command
from django.test import TestCase
from django.urls import NoReverseMatch, reverse

from apps.catalogos.models import NumeroParte
from apps.conceptos.forms import ConceptoForm
from apps.conceptos.models import (
    Concepto,
    DocumentoConceptos,
    HistorialCoincidencia,
    PatronSerie,
)
from apps.conceptos.services import analizar_historial_para_propuestas
from apps.conceptos.views import _buscar_sugerencias_historial


class ConceptosModelTests(TestCase):
    def test_crear_documento_genera_folio(self):
        documento = DocumentoConceptos.objects.create()

        self.assertEqual(documento.folio, 'FACT-000001')

    def test_crear_concepto_calcula_total_concepto(self):
        documento = DocumentoConceptos.objects.create()

        concepto = Concepto.objects.create(
            documento=documento,
            descripcion='Sensor',
            cantidad=Decimal('2.0000'),
            precio_unitario=Decimal('10.500000'),
        )

        self.assertEqual(concepto.total_concepto, Decimal('21.000000'))

    def test_documento_recalcula_total(self):
        documento = DocumentoConceptos.objects.create()

        Concepto.objects.create(
            documento=documento,
            descripcion='Sensor',
            cantidad=Decimal('2.0000'),
            precio_unitario=Decimal('10.000000'),
        )
        Concepto.objects.create(
            documento=documento,
            descripcion='Cable',
            cantidad=Decimal('1.0000'),
            precio_unitario=Decimal('5.000000'),
        )

        documento.refresh_from_db()
        self.assertEqual(documento.total, Decimal('25.000000'))

    def test_concepto_requiere_numero_parte_o_serie_o_descripcion(self):
        form = ConceptoForm(
            data={
                'numero_parte': '',
                'serie': '',
                'modelo': '',
                'descripcion': '',
                'cantidad': '1',
                'precio_unitario': '0',
                'orden': '0',
            }
        )

        self.assertFalse(form.is_valid())
        self.assertIn('Debe capturar numero de parte, serie o descripcion.', form.errors['__all__'])

    def test_cantidad_debe_ser_mayor_a_cero(self):
        form = ConceptoForm(
            data={
                'descripcion': 'Sensor',
                'cantidad': '0',
                'precio_unitario': '0',
                'orden': '0',
            }
        )

        self.assertFalse(form.is_valid())
        self.assertIn('La cantidad debe ser mayor a 0.', form.errors['cantidad'])

    def test_precio_unitario_permite_cero_pero_no_negativo(self):
        form_cero = ConceptoForm(
            data={
                'descripcion': 'Sensor',
                'cantidad': '1',
                'precio_unitario': '0',
                'orden': '0',
            }
        )
        form_negativo = ConceptoForm(
            data={
                'descripcion': 'Sensor',
                'cantidad': '1',
                'precio_unitario': '-1',
                'orden': '0',
            }
        )

        self.assertTrue(form_cero.is_valid())
        self.assertFalse(form_negativo.is_valid())
        self.assertIn(
            'El precio unitario no puede ser negativo.',
            form_negativo.errors['precio_unitario'],
        )

    def test_nuevo_concepto_inicia_con_cantidad_uno(self):
        form = ConceptoForm()

        self.assertEqual(form['cantidad'].value(), '1')

    def test_nuevo_concepto_inicia_con_precio_unitario_cero(self):
        form = ConceptoForm()

        self.assertEqual(form['precio_unitario'].value(), '0')

    def test_edicion_conserva_cantidad_y_precio_existentes(self):
        documento = DocumentoConceptos.objects.create()
        concepto = Concepto.objects.create(
            documento=documento,
            descripcion='Sensor',
            cantidad=Decimal('5.0000'),
            precio_unitario=Decimal('12.340000'),
        )

        form = ConceptoForm(instance=concepto)

        self.assertEqual(form['cantidad'].value(), Decimal('5.0000'))
        self.assertEqual(form['precio_unitario'].value(), Decimal('12.340000'))

    def test_comando_auditar_patrones_muestra_resumen_sin_modificar_datos(self):
        PatronSerie.objects.create(
            prefix='ABC',
            numero_parte='NP-1',
            source='manual',
            sample_size=2,
            confidence=Decimal('1.2000'),
        )
        salida = StringIO()

        call_command('auditar_patrones', stdout=salida)

        contenido = salida.getvalue()
        self.assertIn('Total PatronSerie: 1', contenido)
        self.assertIn('PatronSerie por source', contenido)
        self.assertIn('Prefijos duplicados exactos: 0', contenido)
        self.assertIn('Duplicados logicos restantes: 0', contenido)
        self.assertEqual(PatronSerie.objects.count(), 1)

    def test_modelo_propuesta_patron_serie_ya_no_existe(self):
        from django.apps import apps

        with self.assertRaises(LookupError):
            apps.get_model('conceptos', 'PropuestaPatronSerie')

    def test_guardar_concepto_en_documento_confirmado_registra_historial(self):
        documento = DocumentoConceptos.objects.create(
            status=DocumentoConceptos.STATUS_CONFIRMADO,
        )

        concepto = Concepto.objects.create(
            documento=documento,
            numero_parte='np-001',
            serie=' ser-001 ',
            modelo='mod-a',
            descripcion='Sensor',
            cantidad=Decimal('1'),
            precio_unitario=Decimal('10'),
        )

        historial = HistorialCoincidencia.objects.get(concepto=concepto)
        self.assertEqual(historial.serie, 'SER-001')
        self.assertEqual(historial.numero_parte, 'NP-001')
        self.assertTrue(historial.usar_para_biblioteca)

    def test_documento_cancelado_no_genera_historial_al_guardar_concepto(self):
        documento = DocumentoConceptos.objects.create(
            status=DocumentoConceptos.STATUS_CANCELADO,
        )

        Concepto.objects.create(
            documento=documento,
            numero_parte='NP-CAN',
            serie='SER-CAN',
            descripcion='Cancelado',
            cantidad=Decimal('1'),
            precio_unitario=Decimal('10'),
        )

        self.assertFalse(HistorialCoincidencia.objects.exists())


class PatronSerieAprendizajeTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(
            username='analista',
            password='clave-segura',
        )

    def _evidencia(
        self,
        serie,
        numero_parte='NP-001',
        modelo='MOD-A',
        descripcion='Sensor',
        usar_para_biblioteca=True,
    ):
        return HistorialCoincidencia.objects.create(
            serie=serie,
            numero_parte=numero_parte,
            modelo=modelo,
            descripcion=descripcion,
            firma_json={
                'numero_parte': numero_parte,
                'modelo': modelo,
                'descripcion': descripcion,
            },
            firma_texto=f'numero_parte={numero_parte}|modelo={modelo}|descripcion={descripcion}',
            usar_para_biblioteca=usar_para_biblioteca,
        )

    def test_una_serie_crea_patron_observado(self):
        self._evidencia('OBS001')

        patrones = analizar_historial_para_propuestas()

        self.assertEqual(len(patrones), 1)
        patron = PatronSerie.objects.get()
        self.assertEqual(patron.estado, PatronSerie.ESTADO_OBSERVADO)
        self.assertFalse(patron.activo)
        self.assertEqual(patron.sample_size, 1)
        self.assertEqual(patron.series_unicas, 1)
        self.assertEqual(patron.evidencias_totales, 1)

    def test_dos_series_unicas_crean_patron_en_crecimiento(self):
        self._evidencia('ABC001')
        self._evidencia('ABC002')

        patrones = analizar_historial_para_propuestas()

        self.assertEqual(len(patrones), 1)
        patron = PatronSerie.objects.get()
        self.assertEqual(patron.estado, PatronSerie.ESTADO_EN_CRECIMIENTO)
        self.assertFalse(patron.activo)
        self.assertEqual(patron.sample_size, 2)
        self.assertEqual(patron.series_unicas, 2)

    def test_tres_series_unicas_compatibles_autoaprueban_patron(self):
        self._evidencia('ABC001')
        self._evidencia('ABC002')
        self._evidencia('ABC003')

        patrones = analizar_historial_para_propuestas()

        self.assertEqual(len(patrones), 1)
        patron = PatronSerie.objects.get()
        self.assertEqual(patron.prefix, 'ABC00')
        self.assertEqual(patron.numero_parte, 'NP-001')
        self.assertEqual(patron.sample_size, 3)
        self.assertEqual(patron.series_unicas, 3)
        self.assertEqual(patron.evidencias_totales, 3)
        self.assertEqual(patron.confidence, Decimal('0.6000'))
        self.assertEqual(patron.estado, PatronSerie.ESTADO_APROBADO)
        self.assertTrue(patron.activo)
        self.assertEqual(patron.source, 'evidence_autoapproved')

    def test_duplicados_de_misma_serie_no_aumentan_sample_size(self):
        self._evidencia('ABC001')
        self._evidencia('ABC001')
        self._evidencia('ABC002')

        analizar_historial_para_propuestas()

        patron = PatronSerie.objects.get()
        self.assertEqual(patron.sample_size, 2)
        self.assertEqual(patron.series_unicas, 2)
        self.assertEqual(patron.evidencias_totales, 3)
        self.assertEqual(patron.estado, PatronSerie.ESTADO_EN_CRECIMIENTO)

    def test_prefijo_mas_corto_que_minimo_crea_patron_conflicto(self):
        self._evidencia('AB1001')
        self._evidencia('AB2002')
        self._evidencia('AB3003')

        analizar_historial_para_propuestas()

        patron = PatronSerie.objects.get()
        self.assertEqual(patron.estado, PatronSerie.ESTADO_CONFLICTO)
        self.assertFalse(patron.activo)
        self.assertEqual(patron.motivo_conflicto, 'prefijo demasiado corto')

    def test_conflicto_de_prefijo_con_varios_numero_parte_queda_marcado(self):
        for serie in ('XYZ001', 'XYZ002', 'XYZ003'):
            self._evidencia(serie, numero_parte='NP-001')
        for serie in ('XYZ004', 'XYZ005', 'XYZ006'):
            self._evidencia(serie, numero_parte='NP-002')

        analizar_historial_para_propuestas()

        self.assertTrue(
            PatronSerie.objects.filter(
                estado=PatronSerie.ESTADO_CONFLICTO,
                motivo_conflicto='mismo prefijo asociado a varios numero_parte',
            ).exists()
        )

    def test_colision_de_prefijo_bajo_umbral_queda_en_conflicto(self):
        for serie in ('AB-001', 'AB-002'):
            self._evidencia(serie, numero_parte='NP-A')
        for serie in ('AB-101', 'AB-102'):
            self._evidencia(serie, numero_parte='NP-B')

        analizar_historial_para_propuestas()

        patrones = PatronSerie.objects.order_by('numero_parte')
        self.assertEqual(patrones.count(), 2)
        self.assertTrue(all(patron.estado == PatronSerie.ESTADO_CONFLICTO for patron in patrones))
        self.assertTrue(all(not patron.activo for patron in patrones))

    def test_rutas_operativas_de_propuestas_historicas_retiradas(self):
        with self.assertRaises(NoReverseMatch):
            reverse('conceptos:propuestas_patron_list')
        with self.assertRaises(NoReverseMatch):
            reverse('conceptos:propuesta_patron_aprobar', kwargs={'pk': 1})
        with self.assertRaises(NoReverseMatch):
            reverse('conceptos:propuesta_patron_rechazar', kwargs={'pk': 1})

    def test_no_usa_evidencia_con_usar_para_biblioteca_false(self):
        self._evidencia('NOB001', usar_para_biblioteca=False)
        self._evidencia('NOB002', usar_para_biblioteca=False)
        self._evidencia('NOB003', usar_para_biblioteca=False)

        propuestas = analizar_historial_para_propuestas()

        self.assertEqual(propuestas, [])
        self.assertFalse(PatronSerie.objects.exists())


class ConceptosViewsTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(
            username='operador',
            password='clave-segura',
        )

    def _login(self):
        self.client.force_login(self.user)

    def _grant(self, *codenames):
        permisos = Permission.objects.filter(
            content_type__app_label='conceptos',
            codename__in=codenames,
        )
        self.user.user_permissions.add(*permisos)

    def _crear_documento(self, status=DocumentoConceptos.STATUS_BORRADOR):
        return DocumentoConceptos.objects.create(status=status, usuario=self.user)

    def _crear_concepto(self, documento, numero_parte, orden):
        return Concepto.objects.create(
            documento=documento,
            numero_parte=numero_parte,
            serie=f'SER-{numero_parte}',
            descripcion=f'Descripcion {numero_parte}',
            cantidad=Decimal('1'),
            precio_unitario=Decimal(str(orden)),
            orden=orden,
        )

    def _csv_upload(self, contenido, nombre='conceptos.csv'):
        return SimpleUploadedFile(
            nombre,
            contenido.encode('utf-8-sig'),
            content_type='text/csv',
        )

    def _xlsx_upload(self, filas, nombre='conceptos.xlsx'):
        from openpyxl import Workbook

        workbook = Workbook()
        sheet = workbook.active
        for fila in filas:
            sheet.append(fila)
        buffer = BytesIO()
        workbook.save(buffer)
        buffer.seek(0)
        return SimpleUploadedFile(
            nombre,
            buffer.read(),
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        )

    def _crear_historial_sugerencia(self, documento, serie, usar_para_biblioteca=True):
        return HistorialCoincidencia.objects.create(
            documento=documento,
            serie=serie,
            numero_parte='NP-HIST',
            modelo='MOD-HIST',
            descripcion='Descripcion historial',
            usar_para_biblioteca=usar_para_biblioteca,
        )

    def _buscar_historial(self, serie):
        return _buscar_sugerencias_historial({'numero_parte': '', 'serie': serie})

    def test_sugerencias_no_usan_historial_fuera_de_biblioteca(self):
        documento = self._crear_documento(status=DocumentoConceptos.STATUS_CONFIRMADO)
        historial = self._crear_historial_sugerencia(
            documento, 'SER-FUERA', usar_para_biblioteca=False
        )

        sugerencias = self._buscar_historial('SER-FUERA')

        self.assertFalse(any(item['id'] == historial.pk for item in sugerencias))

    def test_sugerencias_no_usan_historial_de_documento_borrador(self):
        documento = self._crear_documento(status=DocumentoConceptos.STATUS_BORRADOR)
        historial = self._crear_historial_sugerencia(documento, 'SER-BORRADOR')

        sugerencias = self._buscar_historial('SER-BORRADOR')

        self.assertFalse(any(item['id'] == historial.pk for item in sugerencias))

    def test_sugerencias_si_usan_historial_confirmado_de_biblioteca(self):
        documento = self._crear_documento(status=DocumentoConceptos.STATUS_CONFIRMADO)
        historial = self._crear_historial_sugerencia(documento, 'SER-CONFIRMADO')

        sugerencias = self._buscar_historial('SER-CONFIRMADO')

        self.assertEqual([item['id'] for item in sugerencias], [historial.pk])

    def test_listado_requiere_login(self):
        response = self.client.get(reverse('conceptos:documentos_list'))

        self.assertEqual(response.status_code, 302)
        self.assertIn('/accounts/login/', response['Location'])

    def test_listado_requiere_permiso_view(self):
        self._login()

        response = self.client.get(reverse('conceptos:documentos_list'))

        self.assertEqual(response.status_code, 302)
        self.assertIn('/accounts/login/', response['Location'])

    def test_usuario_con_permiso_view_accede_listado(self):
        self._grant('view_documentoconceptos')
        self._login()

        response = self.client.get(reverse('conceptos:documentos_list'))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Facturas')

    def test_listado_filtra_por_folio_y_estado(self):
        primero = self._crear_documento()
        segundo = self._crear_documento(status=DocumentoConceptos.STATUS_CONFIRMADO)
        self._grant('view_documentoconceptos')
        self._login()

        response = self.client.get(
            reverse('conceptos:documentos_list'),
            {'q': primero.folio, 'estado': DocumentoConceptos.STATUS_BORRADOR},
        )

        self.assertContains(response, primero.folio)
        self.assertNotContains(response, segundo.folio)
        self.assertContains(response, 'Filtros activos:')
        self.assertContains(response, 'Limpiar filtros')

    def test_listado_ajax_devuelve_solo_partial_de_resultados(self):
        documento = self._crear_documento()
        self._grant('view_documentoconceptos')
        self._login()

        response = self.client.get(
            reverse('conceptos:documentos_list'),
            {'q': documento.folio},
            HTTP_X_REQUESTED_WITH='XMLHttpRequest',
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, documento.folio)
        self.assertNotContains(response, 'Facturas')
        self.assertNotContains(response, 'data-filtros-documentos')

    def test_listado_paginacion_conserva_filtros(self):
        for _ in range(26):
            self._crear_documento()
        self._grant('view_documentoconceptos')
        self._login()

        response = self.client.get(
            reverse('conceptos:documentos_list'),
            {'q': 'FACT-', 'page': 2},
        )

        self.assertContains(response, 'page=1')
        self.assertContains(response, 'q=FACT-')

    def test_listado_sin_resultados_por_filtro_muestra_estado_util(self):
        self._crear_documento()
        self._grant('view_documentoconceptos', 'add_documentoconceptos')
        self._login()

        response = self.client.get(reverse('conceptos:documentos_list'), {'q': 'FACT-NO-EXISTE'})

        self.assertContains(response, 'No encontramos facturas con los filtros aplicados.')
        self.assertContains(response, 'Limpiar filtros')
        self.assertContains(response, 'Nueva factura')

    def test_listado_vacio_muestra_accion_para_crear(self):
        self._grant('view_documentoconceptos', 'add_documentoconceptos')
        self._login()

        response = self.client.get(reverse('conceptos:documentos_list'))

        self.assertContains(response, 'Aún no hay facturas cargadas.')
        self.assertContains(response, 'Nueva factura')

    def test_detalle_muestra_resumen_compacto_y_acciones_del_borrador(self):
        documento = self._crear_documento()
        self._grant(
            'view_documentoconceptos',
            'change_documentoconceptos',
            'puede_confirmar_documentoconceptos',
            'puede_cancelar_documentoconceptos',
        )
        self._login()

        response = self.client.get(
            reverse('conceptos:documento_detail', kwargs={'pk': documento.pk})
        )

        self.assertContains(response, documento.folio)
        self.assertContains(response, 'Borrador', count=1)
        self.assertContains(response, 'Conceptos')
        self.assertContains(response, '· Total: 0.00')
        self.assertContains(response, 'Fuente: manual')
        self.assertContains(response, '✎ Editar observaciones')
        self.assertContains(response, '⬇ Exportar Word')
        self.assertContains(response, '← Volver')
        self.assertContains(response, '＋ Agregar concepto')
        self.assertContains(response, '⇧ Importar conceptos')
        self.assertContains(response, '✓ Confirmar factura')
        self.assertContains(response, '⚠ Cancelar factura')
        self.assertNotContains(response, 'Sin observaciones.')

    def test_usuario_con_permiso_add_crea_documento(self):
        self._grant('add_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:documento_create'),
            {'observaciones': 'Documento de prueba'},
        )

        self.assertEqual(response.status_code, 302)
        self.assertTrue(DocumentoConceptos.objects.filter(folio='FACT-000001').exists())

    def test_usuario_con_change_agrega_concepto(self):
        documento = self._crear_documento()
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_create', kwargs={'pk': documento.pk}),
            {
                'descripcion': 'Sensor',
                'cantidad': '2',
                'precio_unitario': '10',
                'orden': '0',
            },
        )

        self.assertEqual(response.status_code, 302)
        documento.refresh_from_db()
        self.assertEqual(documento.total, Decimal('20.000000'))

    def test_captura_manual_bloquea_serie_duplicada_en_mismo_documento(self):
        documento = self._crear_documento()
        self._crear_concepto(documento, 'NP-DUP', 1)
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_create', kwargs={'pk': documento.pk}),
            {
                'serie': ' ser-np-dup ',
                'descripcion': 'Duplicado',
                'cantidad': '1',
                'precio_unitario': '0',
                'orden': '0',
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'La serie ya existe en este documento.')
        self.assertEqual(Concepto.objects.filter(documento=documento).count(), 1)

    def test_edicion_no_bloquea_la_misma_serie_del_concepto_actual(self):
        documento = self._crear_documento()
        concepto = self._crear_concepto(documento, 'NP-EDIT', 1)
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse(
                'conceptos:concepto_update',
                kwargs={'pk': documento.pk, 'concepto_id': concepto.pk},
            ),
            {
                'numero_parte': concepto.numero_parte,
                'serie': ' ser-np-edit ',
                'descripcion': 'Descripcion editada',
                'cantidad': '1',
                'precio_unitario': '2',
                'orden': '1',
            },
        )

        self.assertEqual(response.status_code, 302)
        concepto.refresh_from_db()
        self.assertEqual(concepto.serie, 'ser-np-edit')

    def test_misma_serie_en_otro_documento_si_se_permite(self):
        documento_origen = self._crear_documento()
        self._crear_concepto(documento_origen, 'NP-OTRO', 1)
        documento_destino = self._crear_documento()
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_create', kwargs={'pk': documento_destino.pk}),
            {
                'serie': 'SER-NP-OTRO',
                'descripcion': 'Permitido',
                'cantidad': '1',
                'precio_unitario': '0',
                'orden': '0',
            },
        )

        self.assertEqual(response.status_code, 302)
        self.assertEqual(Concepto.objects.filter(serie='SER-NP-OTRO').count(), 2)

    def test_buscar_numero_parte_activo_devuelve_modelo_y_descripcion(self):
        documento = self._crear_documento()
        NumeroParte.objects.create(
            numero_parte='NP-001',
            modelo='MOD-A',
            descripcion='Sensor activo',
            fraccion='',
            activo=True,
        )
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_create', kwargs={'pk': documento.pk}),
            {
                'accion': 'buscar_numero_parte',
                'numero_parte': '  NP-001  ',
                'serie': 'SER-1',
                'cantidad': '2',
                'precio_unitario': '10',
                'orden': '0',
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'value="NP-001"')
        self.assertContains(response, 'value="SER-1"')
        self.assertContains(response, 'value="MOD-A"')
        self.assertContains(response, 'Sensor activo')
        self.assertFalse(Concepto.objects.exists())

    def test_buscar_numero_parte_inexistente_no_bloquea_captura_manual(self):
        documento = self._crear_documento()
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_create', kwargs={'pk': documento.pk}),
            {
                'accion': 'buscar_numero_parte',
                'numero_parte': 'NP-NO',
                'descripcion': 'Manual',
                'cantidad': '1',
                'precio_unitario': '0',
                'orden': '0',
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(
            response,
            'No se encontró número de parte activo; puedes capturar los datos manualmente.',
        )
        self.assertContains(response, 'Lista / Válida')
        self.assertFalse(Concepto.objects.exists())

    def test_numero_parte_inactivo_no_se_usa_como_sugerencia_valida(self):
        documento = self._crear_documento()
        NumeroParte.objects.create(
            numero_parte='NP-INACTIVO',
            modelo='MOD-Z',
            descripcion='Sensor inactivo',
            fraccion='',
            activo=False,
        )
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_create', kwargs={'pk': documento.pk}),
            {
                'accion': 'buscar_numero_parte',
                'numero_parte': 'NP-INACTIVO',
                'cantidad': '1',
                'precio_unitario': '0',
                'orden': '0',
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'El número de parte existe pero está inactivo.')
        self.assertNotContains(response, 'value="MOD-Z"')
        self.assertNotContains(response, 'Sensor inactivo')

    def test_guardar_concepto_con_numero_parte_encontrado_calcula_total(self):
        documento = self._crear_documento()
        NumeroParte.objects.create(
            numero_parte='NP-001',
            modelo='MOD-A',
            descripcion='Sensor activo',
            fraccion='',
            activo=True,
        )
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_create', kwargs={'pk': documento.pk}),
            {
                'accion': 'guardar',
                'numero_parte': 'NP-001',
                'modelo': 'MOD-A',
                'descripcion': 'Sensor activo',
                'cantidad': '2',
                'precio_unitario': '10',
                'orden': '0',
            },
        )

        self.assertEqual(response.status_code, 302)
        concepto = Concepto.objects.get(documento=documento)
        self.assertEqual(concepto.total_concepto, Decimal('20.000000'))

    def test_guardar_concepto_manual_sin_match_sigue_permitido(self):
        documento = self._crear_documento()
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_create', kwargs={'pk': documento.pk}),
            {
                'accion': 'guardar',
                'numero_parte': 'NP-MANUAL',
                'descripcion': 'Captura manual',
                'cantidad': '1',
                'precio_unitario': '0',
                'orden': '0',
            },
        )

        self.assertEqual(response.status_code, 302)
        self.assertTrue(Concepto.objects.filter(numero_parte='NP-MANUAL').exists())

    def test_busqueda_historial_no_usa_conceptos_de_documentos_borrador(self):
        documento = self._crear_documento()
        borrador = self._crear_documento(status=DocumentoConceptos.STATUS_BORRADOR)
        Concepto.objects.create(
            documento=borrador,
            serie='SER-BORRADOR',
            descripcion='Sensor en borrador',
            cantidad=Decimal('1'),
            precio_unitario=Decimal('10'),
        )
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_create', kwargs={'pk': documento.pk}),
            {
                'accion': 'buscar_historial',
                'serie': 'SER-BORRADOR',
                'cantidad': '1',
                'precio_unitario': '0',
                'orden': '0',
            },
        )

        self.assertContains(
            response,
            'No se encontraron coincidencias en historial confirmado.',
        )
        self.assertNotContains(response, 'Sensor en borrador')

    def test_busqueda_historial_no_usa_conceptos_de_documentos_cancelados(self):
        documento = self._crear_documento()
        cancelado = self._crear_documento(status=DocumentoConceptos.STATUS_CANCELADO)
        Concepto.objects.create(
            documento=cancelado,
            serie='SER-CANCELADO',
            descripcion='Sensor cancelado',
            cantidad=Decimal('1'),
            precio_unitario=Decimal('10'),
        )
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_create', kwargs={'pk': documento.pk}),
            {
                'accion': 'buscar_historial',
                'serie': 'SER-CANCELADO',
                'cantidad': '1',
                'precio_unitario': '0',
                'orden': '0',
            },
        )

        self.assertContains(
            response,
            'No se encontraron coincidencias en historial confirmado.',
        )
        self.assertNotContains(response, 'Sensor cancelado')

    def test_busqueda_historial_usa_conceptos_confirmados(self):
        documento = self._crear_documento()
        confirmado = self._crear_documento(status=DocumentoConceptos.STATUS_CONFIRMADO)
        Concepto.objects.create(
            documento=confirmado,
            numero_parte='NP-HIST',
            serie='SER-HIST',
            modelo='MOD-H',
            descripcion='Sensor historico',
            cantidad=Decimal('1'),
            precio_unitario=Decimal('15'),
        )
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_create', kwargs={'pk': documento.pk}),
            {
                'accion': 'buscar_historial',
                'serie': 'SER-HIST',
                'cantidad': '1',
                'precio_unitario': '0',
                'orden': '0',
            },
        )

        self.assertContains(response, 'Selecciona una sugerencia para precargar datos.')
        self.assertContains(response, 'Sensor historico')
        self.assertContains(response, 'Historial')

    def test_busqueda_historial_por_serie_exacta_normalizada(self):
        documento = self._crear_documento()
        confirmado = self._crear_documento(status=DocumentoConceptos.STATUS_CONFIRMADO)
        Concepto.objects.create(
            documento=confirmado,
            serie='SER-NORM',
            descripcion='Sensor normalizado',
            cantidad=Decimal('1'),
            precio_unitario=Decimal('10'),
        )
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_create', kwargs={'pk': documento.pk}),
            {
                'accion': 'buscar_historial',
                'serie': '  ser-norm  ',
                'cantidad': '1',
                'precio_unitario': '0',
                'orden': '0',
            },
        )

        self.assertContains(response, 'Sensor normalizado')

    def test_busqueda_historial_por_numero_parte_exacta_normalizada(self):
        documento = self._crear_documento()
        confirmado = self._crear_documento(status=DocumentoConceptos.STATUS_CONFIRMADO)
        Concepto.objects.create(
            documento=confirmado,
            numero_parte='NP-NORM',
            descripcion='Concepto por numero',
            cantidad=Decimal('1'),
            precio_unitario=Decimal('10'),
        )
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_create', kwargs={'pk': documento.pk}),
            {
                'accion': 'buscar_historial',
                'numero_parte': '  np-norm  ',
                'cantidad': '1',
                'precio_unitario': '0',
                'orden': '0',
            },
        )

        self.assertContains(response, 'Concepto por numero')

    def test_busqueda_historial_por_numero_parte_y_serie(self):
        documento = self._crear_documento()
        confirmado = self._crear_documento(status=DocumentoConceptos.STATUS_CONFIRMADO)
        Concepto.objects.create(
            documento=confirmado,
            numero_parte='NP-BOTH',
            serie='SER-BOTH',
            descripcion='Coincidencia exacta de ambos',
            cantidad=Decimal('1'),
            precio_unitario=Decimal('10'),
        )
        Concepto.objects.create(
            documento=confirmado,
            numero_parte='NP-BOTH',
            serie='SER-OTRA',
            descripcion='Solo coincide numero',
            cantidad=Decimal('1'),
            precio_unitario=Decimal('10'),
        )
        Concepto.objects.create(
            documento=confirmado,
            numero_parte='NP-OTRO',
            serie='SER-BOTH',
            descripcion='Solo coincide serie',
            cantidad=Decimal('1'),
            precio_unitario=Decimal('10'),
        )
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_create', kwargs={'pk': documento.pk}),
            {
                'accion': 'buscar_historial',
                'numero_parte': 'np-both',
                'serie': 'ser-both',
                'cantidad': '1',
                'precio_unitario': '0',
                'orden': '0',
            },
        )

        self.assertContains(response, 'Coincidencia exacta de ambos')
        self.assertNotContains(response, 'Solo coincide numero')
        self.assertNotContains(response, 'Solo coincide serie')

    def test_busqueda_historial_sin_numero_ni_serie_muestra_mensaje_claro(self):
        documento = self._crear_documento()
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_create', kwargs={'pk': documento.pk}),
            {
                'accion': 'buscar_historial',
                'descripcion': 'temperatura',
                'cantidad': '1',
                'precio_unitario': '0',
                'orden': '0',
            },
        )

        self.assertContains(response, 'Captura número de parte o serie para buscar historial.')

    def test_usar_sugerencia_precarga_datos_sin_reemplazar_cantidad(self):
        documento = self._crear_documento()
        confirmado = self._crear_documento(status=DocumentoConceptos.STATUS_CONFIRMADO)
        concepto_historial = Concepto.objects.create(
            documento=confirmado,
            numero_parte='NP-SUG',
            serie='SER-SUG',
            modelo='MOD-SUG',
            descripcion='Descripcion sugerida',
            cantidad=Decimal('1'),
            precio_unitario=Decimal('99.500000'),
        )
        sugerencia = HistorialCoincidencia.objects.get(concepto=concepto_historial)
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_create', kwargs={'pk': documento.pk}),
            {
                'accion': 'usar_sugerencia',
                'sugerencia_id': sugerencia.pk,
                'serie': 'SER-CAPTURADA',
                'cantidad': '7',
                'precio_unitario': '11',
                'orden': '0',
            },
        )

        self.assertContains(response, 'value="NP-SUG"')
        self.assertContains(response, 'value="SER-CAPTURADA"')
        self.assertContains(response, 'value="MOD-SUG"')
        self.assertContains(response, 'Descripcion sugerida')
        self.assertContains(response, 'value="7"')
        self.assertContains(response, 'value="11"')
        self.assertFalse(Concepto.objects.filter(documento=documento).exists())

    def test_buscar_historial_no_consulta_catalogo_automaticamente(self):
        documento = self._crear_documento()
        NumeroParte.objects.create(
            numero_parte='NP-ACTIVO',
            modelo='MOD-ACTIVO',
            descripcion='Descripcion activa',
            fraccion='',
            activo=True,
        )
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_create', kwargs={'pk': documento.pk}),
            {
                'accion': 'buscar_historial',
                'numero_parte': 'NP-ACTIVO',
                'cantidad': '3',
                'precio_unitario': '12',
                'orden': '0',
            },
        )

        self.assertContains(response, 'No se encontraron coincidencias en historial confirmado.')
        self.assertNotContains(response, 'value="MOD-ACTIVO"')
        self.assertNotContains(response, 'Descripcion activa')
        self.assertFalse(Concepto.objects.filter(documento=documento).exists())

    def test_lupa_numero_parte_no_consulta_historial_automaticamente(self):
        documento = self._crear_documento()
        confirmado = self._crear_documento(status=DocumentoConceptos.STATUS_CONFIRMADO)
        Concepto.objects.create(
            documento=confirmado,
            numero_parte='NP-HIST-AUTO',
            descripcion='Descripcion solo en historial',
            cantidad=Decimal('1'),
            precio_unitario=Decimal('10'),
        )
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_create', kwargs={'pk': documento.pk}),
            {
                'accion': 'buscar_numero_parte',
                'numero_parte': 'NP-HIST-AUTO',
                'cantidad': '1',
                'precio_unitario': '0',
                'orden': '0',
            },
        )

        self.assertContains(
            response,
            'No se encontró número de parte activo; puedes capturar los datos manualmente.',
        )
        self.assertNotContains(response, 'Descripcion solo en historial')

    def test_lupa_serie_no_consulta_historial_automaticamente(self):
        documento = self._crear_documento()
        confirmado = self._crear_documento(status=DocumentoConceptos.STATUS_CONFIRMADO)
        Concepto.objects.create(
            documento=confirmado,
            serie='SER-SOLO-HIST',
            descripcion='Historial por serie',
            cantidad=Decimal('1'),
            precio_unitario=Decimal('10'),
        )
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_create', kwargs={'pk': documento.pk}),
            {
                'accion': 'buscar_serie',
                'serie': 'SER-SOLO-HIST',
                'cantidad': '1',
                'precio_unitario': '0',
                'orden': '0',
            },
        )

        self.assertContains(response, 'No se encontraron patrones activos para la serie.')
        self.assertNotContains(response, 'Historial por serie')

    def test_patron_serie_activo_por_prefijo_sugiere_datos(self):
        documento = self._crear_documento()
        PatronSerie.objects.create(
            prefix='ps-',
            numero_parte='NP-PATRON',
            modelo='MOD-P',
            descripcion='Descripcion patron',
            sample_size=4,
            confidence=Decimal('0.7500'),
        )
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_create', kwargs={'pk': documento.pk}),
            {
                'accion': 'buscar_serie',
                'serie': 'PS-12345',
                'cantidad': '1',
                'precio_unitario': '0',
                'orden': '0',
            },
        )

        self.assertContains(response, 'NP-PATRON')
        self.assertContains(response, 'Descripcion patron')
        self.assertContains(response, 'Patron')
        self.assertContains(response, 'Prefijo: PS-')
        self.assertContains(response, 'Muestra: 4')

    def test_lupa_numero_parte_usa_patron_si_no_hay_catalogo(self):
        documento = self._crear_documento()
        PatronSerie.objects.create(
            prefix='NP-P',
            numero_parte='NP-PATRON-NUM',
            modelo='MOD-NUM',
            descripcion='Descripcion patron por numero',
            sample_size=4,
            confidence=Decimal('0.7500'),
        )
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_create', kwargs={'pk': documento.pk}),
            {
                'accion': 'buscar_numero_parte',
                'numero_parte': ' np-patron-num ',
                'serie': 'SER-CAPTURADA',
                'cantidad': '4',
                'precio_unitario': '20',
                'orden': '0',
            },
        )

        self.assertContains(response, 'NP-PATRON-NUM')
        self.assertContains(response, 'Descripcion patron por numero')
        self.assertContains(response, 'Patron')
        self.assertContains(response, 'value="SER-CAPTURADA"')
        self.assertContains(response, 'value="4"')
        self.assertContains(response, 'value="20"')

    def test_lupa_numero_parte_muestra_varios_patrones(self):
        documento = self._crear_documento()
        PatronSerie.objects.create(
            prefix='NP-A',
            numero_parte='NP-MULTI',
            descripcion='Patron multi A',
        )
        PatronSerie.objects.create(
            prefix='NP-B',
            numero_parte='NP-MULTI',
            descripcion='Patron multi B',
        )
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_create', kwargs={'pk': documento.pk}),
            {
                'accion': 'buscar_numero_parte',
                'numero_parte': 'NP-MULTI',
                'cantidad': '1',
                'precio_unitario': '0',
                'orden': '0',
            },
        )

        self.assertContains(response, 'Patron multi A')
        self.assertContains(response, 'Patron multi B')

    def test_patron_serie_inactivo_no_sugiere(self):
        documento = self._crear_documento()
        PatronSerie.objects.create(
            prefix='PX-',
            numero_parte='NP-INACTIVO',
            descripcion='Patron inactivo',
            activo=False,
        )
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_create', kwargs={'pk': documento.pk}),
            {
                'accion': 'buscar_serie',
                'serie': 'PX-123',
                'cantidad': '1',
                'precio_unitario': '0',
                'orden': '0',
            },
        )

        self.assertContains(
            response,
            'No se encontraron patrones activos para la serie.',
        )
        self.assertNotContains(response, 'Patron inactivo')

    def test_patron_serie_gana_prefijo_mas_largo(self):
        documento = self._crear_documento()
        PatronSerie.objects.create(
            prefix='ABC',
            numero_parte='NP-CORTO',
            descripcion='Patron corto',
        )
        PatronSerie.objects.create(
            prefix='ABC123',
            numero_parte='NP-LARGO',
            descripcion='Patron largo',
        )
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_create', kwargs={'pk': documento.pk}),
            {
                'accion': 'buscar_serie',
                'serie': 'ABC123XYZ',
                'cantidad': '1',
                'precio_unitario': '0',
                'orden': '0',
            },
        )

        self.assertContains(response, 'NP-LARGO')
        self.assertContains(response, 'Patron largo')
        self.assertNotContains(response, 'NP-CORTO')

    def test_patron_serie_empate_de_longitud_ordena_por_prefix(self):
        documento = self._crear_documento()
        PatronSerie.objects.create(
            prefix='EMP',
            numero_parte='NP-B',
            descripcion='Patron empate B',
        )
        PatronSerie.objects.create(
            prefix='EMP',
            numero_parte='NP-A',
            descripcion='Patron empate A',
        )
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_create', kwargs={'pk': documento.pk}),
            {
                'accion': 'buscar_serie',
                'serie': 'EMP-001',
                'cantidad': '1',
                'precio_unitario': '0',
                'orden': '0',
            },
        )

        self.assertContains(response, 'Patron')
        self.assertContains(response, 'Prefijo: EMP')

    def test_usar_patron_no_reemplaza_cantidad_precio_ni_serie(self):
        documento = self._crear_documento()
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_create', kwargs={'pk': documento.pk}),
            {
                'accion': 'usar_sugerencia',
                'sugerencia_origen': 'patron',
                'sugerencia_numero_parte': 'NP-PATRON',
                'sugerencia_modelo': 'MOD-P',
                'sugerencia_descripcion': 'Descripcion patron',
                'serie': 'SER-CAPTURADA',
                'cantidad': '8',
                'precio_unitario': '12.340000',
                'orden': '0',
            },
        )

        self.assertContains(response, 'value="NP-PATRON"')
        self.assertContains(response, 'value="MOD-P"')
        self.assertContains(response, 'Descripcion patron')
        self.assertContains(response, 'value="SER-CAPTURADA"')
        self.assertContains(response, 'value="8"')
        self.assertContains(response, 'value="12.340000"')
        self.assertFalse(Concepto.objects.filter(documento=documento).exists())

    def test_exacto_numero_parte_activo_tiene_prioridad_sobre_patron(self):
        documento = self._crear_documento()
        NumeroParte.objects.create(
            numero_parte='NP-EXACTO',
            modelo='MOD-EX',
            descripcion='Descripcion exacta',
            fraccion='',
            activo=True,
        )
        PatronSerie.objects.create(
            prefix='SER-',
            numero_parte='NP-PATRON',
            descripcion='Descripcion patron',
        )
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_create', kwargs={'pk': documento.pk}),
            {
                'accion': 'buscar_numero_parte',
                'numero_parte': 'NP-EXACTO',
                'serie': 'SER-123',
                'cantidad': '1',
                'precio_unitario': '0',
                'orden': '0',
            },
        )

        self.assertContains(response, 'Descripcion exacta')
        self.assertNotContains(response, 'Patron')
        self.assertNotContains(response, 'Descripcion patron')

    def test_historial_confirmado_tiene_prioridad_sobre_patron(self):
        documento = self._crear_documento()
        confirmado = self._crear_documento(status=DocumentoConceptos.STATUS_CONFIRMADO)
        Concepto.objects.create(
            documento=confirmado,
            serie='SER-HIST-PATRON',
            descripcion='Descripcion historial prioritaria',
            cantidad=Decimal('1'),
            precio_unitario=Decimal('10'),
        )
        PatronSerie.objects.create(
            prefix='SER-HIST',
            numero_parte='NP-PATRON',
            descripcion='Descripcion patron secundaria',
        )
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_create', kwargs={'pk': documento.pk}),
            {
                'accion': 'buscar_historial',
                'serie': 'SER-HIST-PATRON',
                'cantidad': '1',
                'precio_unitario': '0',
                'orden': '0',
            },
        )

        self.assertContains(response, 'Descripcion historial prioritaria')
        self.assertContains(response, 'Historial')
        self.assertNotContains(response, 'Patron')
        self.assertNotContains(response, 'Descripcion patron secundaria')

    def test_patron_no_se_usa_si_no_hay_serie(self):
        documento = self._crear_documento()
        PatronSerie.objects.create(
            prefix='SIN-SERIE',
            numero_parte='NP-PATRON',
            descripcion='Patron sin serie',
        )
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_create', kwargs={'pk': documento.pk}),
            {
                'accion': 'buscar_historial',
                'descripcion': 'SIN-SERIE',
                'cantidad': '1',
                'precio_unitario': '0',
                'orden': '0',
            },
        )

        self.assertContains(
            response,
            'Captura número de parte o serie para buscar historial.',
        )
        self.assertNotContains(response, 'Patron sin serie')

    def test_importacion_no_permite_documento_confirmado(self):
        documento = self._crear_documento(status=DocumentoConceptos.STATUS_CONFIRMADO)
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.get(reverse('conceptos:conceptos_importar', kwargs={'pk': documento.pk}))

        self.assertEqual(response.status_code, 302)
        self.assertIn(reverse('conceptos:documento_detail', kwargs={'pk': documento.pk}), response['Location'])

    def test_importacion_no_permite_documento_cancelado(self):
        documento = self._crear_documento(status=DocumentoConceptos.STATUS_CANCELADO)
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.get(reverse('conceptos:conceptos_importar', kwargs={'pk': documento.pk}))

        self.assertEqual(response.status_code, 302)
        self.assertIn(reverse('conceptos:documento_detail', kwargs={'pk': documento.pk}), response['Location'])

    def test_importacion_csv_valido_genera_preview_sin_guardar(self):
        documento = self._crear_documento()
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:conceptos_importar', kwargs={'pk': documento.pk}),
            {
                'archivo': self._csv_upload(
                    'numero_parte,serie,modelo,descripcion,cantidad,precio_unitario\n'
                    'NP-CSV,SER-CSV,MOD-CSV,Descripcion CSV,2,5\n'
                )
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Preview')
        self.assertContains(response, 'Manual')
        self.assertNotContains(response, 'Usar filas confirmadas para alimentar biblioteca')
        self.assertFalse(Concepto.objects.filter(documento=documento).exists())

    def test_importacion_xlsx_valido_genera_preview(self):
        documento = self._crear_documento()
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:conceptos_importar', kwargs={'pk': documento.pk}),
            {
                'archivo': self._xlsx_upload([
                    ['numero_parte', 'serie', 'modelo', 'descripcion', 'cantidad', 'precio_unitario'],
                    ['NP-XLSX', 'SER-XLSX', 'MOD-X', 'Descripcion XLSX', 1, 3],
                ])
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'NP-XLSX')
        self.assertContains(response, 'Descripcion XLSX')

    def test_cancelar_importacion_limpia_preview_sin_crear_conceptos(self):
        documento = self._crear_documento()
        self._grant('change_documentoconceptos')
        self._login()
        importar_url = reverse('conceptos:conceptos_importar', kwargs={'pk': documento.pk})
        cancelar_url = reverse('conceptos:conceptos_importar_cancelar', kwargs={'pk': documento.pk})

        self.client.post(
            importar_url,
            {'archivo': self._csv_upload(
                'numero_parte,serie,modelo,descripcion,cantidad,precio_unitario\n'
                'NP-CANCEL,SER-CANCEL,MOD,Descripcion,1,0\n'
            )},
        )
        self.assertIn(f'conceptos_importacion_{documento.pk}', self.client.session)

        response = self.client.post(cancelar_url)

        self.assertEqual(response.status_code, 302)
        self.assertEqual(response['Location'], importar_url)
        self.assertNotIn(f'conceptos_importacion_{documento.pk}', self.client.session)
        self.assertFalse(Concepto.objects.filter(documento=documento).exists())
        self.assertFalse(HistorialCoincidencia.objects.exists())

        response = self.client.get(importar_url)
        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, 'id="preview-resumen"')
        self.assertContains(response, 'Importación cancelada. No se aplicaron cambios.')

    def test_cancelar_importacion_requiere_post_y_permiso(self):
        documento = self._crear_documento()
        self._login()
        url = reverse('conceptos:conceptos_importar_cancelar', kwargs={'pk': documento.pk})

        response = self.client.get(url)
        self.assertEqual(response.status_code, 405)

        response = self.client.post(url)
        self.assertEqual(response.status_code, 302)
        self.assertIn('/accounts/login/', response['Location'])

    def test_importacion_numero_parte_activo_autollena_modelo_descripcion(self):
        documento = self._crear_documento()
        NumeroParte.objects.create(
            numero_parte='NP-ACTIVO',
            modelo='MOD-ACT',
            descripcion='Descripcion activa',
            fraccion='',
            activo=True,
        )
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:conceptos_importar', kwargs={'pk': documento.pk}),
            {
                'archivo': self._csv_upload(
                    'numero_parte,serie,modelo,descripcion,cantidad,precio_unitario\n'
                    'NP-ACTIVO,SER-ACT,,,1,0\n'
                )
            },
        )

        self.assertContains(response, 'OK exacto')
        self.assertContains(response, 'MOD-ACT')
        self.assertContains(response, 'Descripcion activa')

    def test_importacion_historial_confirmado_sugiere_datos(self):
        documento = self._crear_documento()
        confirmado = self._crear_documento(status=DocumentoConceptos.STATUS_CONFIRMADO)
        Concepto.objects.create(
            documento=confirmado,
            numero_parte='NP-HIST',
            serie='SER-HIST',
            modelo='MOD-HIST',
            descripcion='Descripcion historial',
            cantidad=Decimal('1'),
            precio_unitario=Decimal('10'),
        )
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:conceptos_importar', kwargs={'pk': documento.pk}),
            {
                'archivo': self._csv_upload(
                    'numero_parte,serie,modelo,descripcion,cantidad,precio_unitario\n'
                    ',SER-HIST,,,1,0\n'
                )
            },
        )

        self.assertContains(response, 'Sugerido por historial')
        self.assertContains(response, 'NP-HIST')
        self.assertContains(response, 'Descripcion historial')

    def test_importacion_patron_serie_sugiere_datos(self):
        documento = self._crear_documento()
        PatronSerie.objects.create(
            prefix='PX-',
            numero_parte='NP-PATRON',
            modelo='MOD-P',
            descripcion='Descripcion patron',
        )
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:conceptos_importar', kwargs={'pk': documento.pk}),
            {
                'archivo': self._csv_upload(
                    'numero_parte,serie,modelo,descripcion,cantidad,precio_unitario\n'
                    ',PX-123,,,1,0\n'
                )
            },
        )

        self.assertContains(response, 'Sugerido por patron')
        self.assertContains(response, 'NP-PATRON')
        self.assertContains(response, 'Descripcion patron')

    def test_importacion_defaults_cantidad_y_precio(self):
        documento = self._crear_documento()
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:conceptos_importar', kwargs={'pk': documento.pk}),
            {
                'archivo': self._csv_upload(
                    'numero_parte,serie,modelo,descripcion,cantidad,precio_unitario\n'
                    'NP-DEF,SER-DEF,,Descripcion defaults,,\n'
                )
            },
        )

        self.assertContains(response, '<td>1</td>')
        self.assertContains(response, '<td>0</td>')

    def test_importacion_cantidad_invalida_marca_error(self):
        documento = self._crear_documento()
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:conceptos_importar', kwargs={'pk': documento.pk}),
            {
                'archivo': self._csv_upload(
                    'numero_parte,serie,modelo,descripcion,cantidad,precio_unitario\n'
                    'NP-ERR,SER-ERR,,Descripcion,error,0\n'
                )
            },
        )

        self.assertContains(response, 'Error')
        self.assertContains(response, 'Cantidad invalida.')

    def test_importacion_precio_negativo_marca_error(self):
        documento = self._crear_documento()
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:conceptos_importar', kwargs={'pk': documento.pk}),
            {
                'archivo': self._csv_upload(
                    'numero_parte,serie,modelo,descripcion,cantidad,precio_unitario\n'
                    'NP-ERR,SER-ERR,,Descripcion,1,-1\n'
                )
            },
        )

        self.assertContains(response, 'Error')
        self.assertContains(response, 'Precio unitario invalido.')

    def test_importacion_marca_duplicada_si_serie_existe_en_documento_destino(self):
        documento = self._crear_documento()
        self._crear_concepto(documento, 'NP-EXISTE', 1)
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:conceptos_importar', kwargs={'pk': documento.pk}),
            {
                'archivo': self._csv_upload(
                    'numero_parte,serie,modelo,descripcion,cantidad,precio_unitario\n'
                    'NP-NUEVO, ser-np-existe ,,Duplicado,1,0\n'
                )
            },
        )

        self.assertContains(response, 'Duplicada')
        self.assertContains(response, 'La serie ya existe en este documento.')
        self.client.post(
            reverse('conceptos:conceptos_importar_confirmar', kwargs={'pk': documento.pk}),
            {'usar_para_biblioteca': '1'},
        )
        self.assertEqual(Concepto.objects.filter(documento=documento).count(), 1)
        self.assertFalse(HistorialCoincidencia.objects.exists())

    def test_importacion_marca_duplicada_a_segunda_aparicion_del_archivo(self):
        documento = self._crear_documento()
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:conceptos_importar', kwargs={'pk': documento.pk}),
            {
                'archivo': self._csv_upload(
                    'numero_parte,serie,modelo,descripcion,cantidad,precio_unitario\n'
                    'NP-1,SER-REP,,Primera,1,0\n'
                    'NP-2, ser-rep ,,Segunda,1,0\n'
                )
            },
        )

        self.assertContains(response, 'Válidas: 1')
        self.assertContains(response, 'Duplicadas: 1')
        self.assertContains(response, 'La serie está repetida dentro del archivo.')

    def test_importacion_duplicadas_no_crean_concepto_ni_historial(self):
        documento = self._crear_documento()
        self._grant('change_documentoconceptos')
        self._login()

        self.client.post(
            reverse('conceptos:conceptos_importar', kwargs={'pk': documento.pk}),
            {
                'archivo': self._csv_upload(
                    'numero_parte,serie,modelo,descripcion,cantidad,precio_unitario\n'
                    'NP-1,SER-DUP,,Primera,1,0\n'
                    'NP-2,SER-DUP,,Segunda,1,0\n'
                )
            },
        )
        self.client.post(
            reverse('conceptos:conceptos_importar_confirmar', kwargs={'pk': documento.pk}),
            {'usar_para_biblioteca': '1'},
        )

        self.assertEqual(Concepto.objects.filter(documento=documento).count(), 1)
        self.assertEqual(Concepto.objects.get(documento=documento).numero_parte, 'NP-1')
        self.assertFalse(HistorialCoincidencia.objects.exists())

    def test_confirmar_importacion_crea_solo_filas_validas_y_recalcula_total(self):
        documento = self._crear_documento()
        self._grant('change_documentoconceptos')
        self._login()

        self.client.post(
            reverse('conceptos:conceptos_importar', kwargs={'pk': documento.pk}),
            {
                'archivo': self._csv_upload(
                    'numero_parte,serie,modelo,descripcion,cantidad,precio_unitario\n'
                    'NP-OK,SER-OK,,Valida,2,5\n'
                    ',,,,1,1\n'
                    'NP-ERR,SER-ERR,,Error,mal,1\n'
                )
            },
        )
        response = self.client.post(
            reverse('conceptos:conceptos_importar_confirmar', kwargs={'pk': documento.pk})
        )

        self.assertEqual(response.status_code, 302)
        self.assertEqual(Concepto.objects.filter(documento=documento).count(), 1)
        documento.refresh_from_db()
        self.assertEqual(documento.total, Decimal('10.000000'))

    def test_importacion_crea_conceptos_sin_historial_mientras_factura_sigue_borrador(self):
        documento = self._crear_documento()
        NumeroParte.objects.create(
            numero_parte='NP-EVID',
            modelo='MOD-EVID',
            descripcion='Descripcion evidencia',
            fraccion='',
            activo=True,
        )
        self._grant('change_documentoconceptos')
        self._login()

        self.client.post(
            reverse('conceptos:conceptos_importar', kwargs={'pk': documento.pk}),
            {
                'archivo': self._csv_upload(
                    'numero_parte,serie,modelo,descripcion,cantidad,precio_unitario\n'
                    'NP-EVID,SER-EVID,,,1,0\n'
                )
            },
        )
        self.client.post(
            reverse('conceptos:conceptos_importar_confirmar', kwargs={'pk': documento.pk}),
            {'usar_para_biblioteca': '1'},
        )

        self.assertEqual(Concepto.objects.filter(documento=documento).count(), 1)
        self.assertEqual(documento.status, DocumentoConceptos.STATUS_BORRADOR)
        self.assertFalse(HistorialCoincidencia.objects.exists())

    def test_importacion_con_opcion_desactivada_no_crea_historial(self):
        documento = self._crear_documento()
        self._grant('change_documentoconceptos')
        self._login()

        self.client.post(
            reverse('conceptos:conceptos_importar', kwargs={'pk': documento.pk}),
            {
                'archivo': self._csv_upload(
                    'numero_parte,serie,modelo,descripcion,cantidad,precio_unitario\n'
                    'NP-SIN,SER-SIN,,Sin evidencia,1,0\n'
                )
            },
        )
        self.client.post(reverse('conceptos:conceptos_importar_confirmar', kwargs={'pk': documento.pk}))

        self.assertEqual(Concepto.objects.filter(documento=documento).count(), 1)
        self.assertFalse(HistorialCoincidencia.objects.exists())

    def test_importacion_con_error_no_crea_evidencia_para_fila_invalida(self):
        documento = self._crear_documento()
        self._grant('change_documentoconceptos')
        self._login()

        self.client.post(
            reverse('conceptos:conceptos_importar', kwargs={'pk': documento.pk}),
            {
                'archivo': self._csv_upload(
                    'numero_parte,serie,modelo,descripcion,cantidad,precio_unitario\n'
                    'NP-OK,SER-OK,,Valida,1,0\n'
                    'NP-ERR,SER-ERR,,Error,mal,0\n'
                )
            },
        )
        self.client.post(
            reverse('conceptos:conceptos_importar_confirmar', kwargs={'pk': documento.pk}),
            {'usar_para_biblioteca': '1'},
        )

        self.assertEqual(Concepto.objects.filter(documento=documento).count(), 1)
        self.assertFalse(HistorialCoincidencia.objects.exists())

    def test_reintento_confirmacion_importacion_no_duplica_historial(self):
        documento = self._crear_documento()
        self._grant('change_documentoconceptos')
        self._login()

        self.client.post(
            reverse('conceptos:conceptos_importar', kwargs={'pk': documento.pk}),
            {
                'archivo': self._csv_upload(
                    'numero_parte,serie,modelo,descripcion,cantidad,precio_unitario\n'
                    'NP-RE,SER-RE,,Reintento,1,0\n'
                )
            },
        )
        url = reverse('conceptos:conceptos_importar_confirmar', kwargs={'pk': documento.pk})
        self.client.post(url)
        self.client.post(url)

        self.assertEqual(Concepto.objects.filter(documento=documento).count(), 1)
        self.assertFalse(HistorialCoincidencia.objects.exists())

    def test_subir_intercambia_orden_con_anterior(self):
        documento = self._crear_documento()
        primero = self._crear_concepto(documento, 'NP-1', 1)
        segundo = self._crear_concepto(documento, 'NP-2', 2)
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_subir', kwargs={'pk': documento.pk, 'concepto_id': segundo.pk})
        )

        self.assertEqual(response.status_code, 302)
        primero.refresh_from_db()
        segundo.refresh_from_db()
        self.assertEqual(primero.orden, 2)
        self.assertEqual(segundo.orden, 1)

    def test_bajar_intercambia_orden_con_siguiente(self):
        documento = self._crear_documento()
        primero = self._crear_concepto(documento, 'NP-1', 1)
        segundo = self._crear_concepto(documento, 'NP-2', 2)
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_bajar', kwargs={'pk': documento.pk, 'concepto_id': primero.pk})
        )

        self.assertEqual(response.status_code, 302)
        primero.refresh_from_db()
        segundo.refresh_from_db()
        self.assertEqual(primero.orden, 2)
        self.assertEqual(segundo.orden, 1)

    def test_ajax_subir_devuelve_json_ok_y_cambia_orden(self):
        documento = self._crear_documento()
        primero = self._crear_concepto(documento, 'NP-1', 1)
        segundo = self._crear_concepto(documento, 'NP-2', 2)
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_subir', kwargs={'pk': documento.pk, 'concepto_id': segundo.pk}),
            HTTP_X_REQUESTED_WITH='XMLHttpRequest',
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.json(),
            {'ok': True, 'concepto_id': segundo.pk, 'accion': 'subir'},
        )
        primero.refresh_from_db()
        segundo.refresh_from_db()
        self.assertEqual(primero.orden, 2)
        self.assertEqual(segundo.orden, 1)

    def test_ajax_bajar_devuelve_json_ok_y_cambia_orden(self):
        documento = self._crear_documento()
        primero = self._crear_concepto(documento, 'NP-1', 1)
        segundo = self._crear_concepto(documento, 'NP-2', 2)
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_bajar', kwargs={'pk': documento.pk, 'concepto_id': primero.pk}),
            HTTP_X_REQUESTED_WITH='XMLHttpRequest',
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.json(),
            {'ok': True, 'concepto_id': primero.pk, 'accion': 'bajar'},
        )
        primero.refresh_from_db()
        segundo.refresh_from_db()
        self.assertEqual(primero.orden, 2)
        self.assertEqual(segundo.orden, 1)

    def test_ajax_reordenar_guarda_orden_completo(self):
        documento = self._crear_documento()
        primero = self._crear_concepto(documento, 'NP-1', 1)
        segundo = self._crear_concepto(documento, 'NP-2', 2)
        tercero = self._crear_concepto(documento, 'NP-3', 3)
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:conceptos_reordenar', kwargs={'pk': documento.pk}),
            data=json.dumps({'orden': [tercero.pk, primero.pk, segundo.pk]}),
            content_type='application/json',
            HTTP_X_REQUESTED_WITH='XMLHttpRequest',
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json(), {'ok': True, 'orden': [tercero.pk, primero.pk, segundo.pk]})
        tercero.refresh_from_db()
        primero.refresh_from_db()
        segundo.refresh_from_db()
        self.assertEqual([tercero.orden, primero.orden, segundo.orden], [1, 2, 3])

    def test_ajax_reordenar_rechaza_concepto_de_otra_factura(self):
        documento = self._crear_documento()
        otro_documento = self._crear_documento()
        concepto = self._crear_concepto(documento, 'NP-1', 1)
        concepto_otro = self._crear_concepto(otro_documento, 'NP-2', 1)
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:conceptos_reordenar', kwargs={'pk': otro_documento.pk}),
            data=json.dumps({'orden': [concepto.pk]}),
            content_type='application/json',
            HTTP_X_REQUESTED_WITH='XMLHttpRequest',
        )

        self.assertEqual(response.status_code, 400)
        self.assertFalse(response.json()['ok'])
        concepto.refresh_from_db()
        concepto_otro.refresh_from_db()
        self.assertEqual(concepto.orden, 1)
        self.assertEqual(concepto_otro.orden, 1)

    def test_ajax_reordenar_rechaza_factura_no_borrador(self):
        documento = self._crear_documento(status=DocumentoConceptos.STATUS_CONFIRMADO)
        concepto = self._crear_concepto(documento, 'NP-1', 1)
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:conceptos_reordenar', kwargs={'pk': documento.pk}),
            data=json.dumps({'orden': [concepto.pk]}),
            content_type='application/json',
        )

        self.assertEqual(response.status_code, 400)
        self.assertFalse(response.json()['ok'])

    def test_ajax_reordenar_rechaza_factura_cancelada(self):
        documento = self._crear_documento(status=DocumentoConceptos.STATUS_CANCELADO)
        concepto = self._crear_concepto(documento, 'NP-1', 1)
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:conceptos_reordenar', kwargs={'pk': documento.pk}),
            data=json.dumps({'orden': [concepto.pk]}),
            content_type='application/json',
        )

        self.assertEqual(response.status_code, 400)
        self.assertFalse(response.json()['ok'])

    def test_ajax_reordenar_respeta_permiso(self):
        documento = self._crear_documento()
        concepto = self._crear_concepto(documento, 'NP-1', 1)
        self._login()

        response = self.client.post(
            reverse('conceptos:conceptos_reordenar', kwargs={'pk': documento.pk}),
            data=json.dumps({'orden': [concepto.pk]}),
            content_type='application/json',
        )

        self.assertEqual(response.status_code, 403)

    def test_primero_no_puede_subir(self):
        documento = self._crear_documento()
        primero = self._crear_concepto(documento, 'NP-1', 1)
        self._crear_concepto(documento, 'NP-2', 2)
        self._grant('change_documentoconceptos')
        self._login()

        self.client.post(
            reverse('conceptos:concepto_subir', kwargs={'pk': documento.pk, 'concepto_id': primero.pk})
        )

        primero.refresh_from_db()
        self.assertEqual(primero.orden, 1)

    def test_ultimo_no_puede_bajar(self):
        documento = self._crear_documento()
        self._crear_concepto(documento, 'NP-1', 1)
        ultimo = self._crear_concepto(documento, 'NP-2', 2)
        self._grant('change_documentoconceptos')
        self._login()

        self.client.post(
            reverse('conceptos:concepto_bajar', kwargs={'pk': documento.pk, 'concepto_id': ultimo.pk})
        )

        ultimo.refresh_from_db()
        self.assertEqual(ultimo.orden, 2)

    def test_documento_confirmado_permite_reordenar(self):
        documento = self._crear_documento(status=DocumentoConceptos.STATUS_CONFIRMADO)
        primero = self._crear_concepto(documento, 'NP-1', 1)
        segundo = self._crear_concepto(documento, 'NP-2', 2)
        historial_inicial = HistorialCoincidencia.objects.count()
        self._grant('change_documentoconceptos')
        self._login()

        self.client.post(
            reverse('conceptos:concepto_subir', kwargs={'pk': documento.pk, 'concepto_id': segundo.pk})
        )

        primero.refresh_from_db()
        segundo.refresh_from_db()
        self.assertEqual(primero.orden, 2)
        self.assertEqual(segundo.orden, 1)
        self.assertEqual(HistorialCoincidencia.objects.count(), historial_inicial)

    def test_documento_cancelado_no_permite_reordenar(self):
        documento = self._crear_documento(status=DocumentoConceptos.STATUS_CANCELADO)
        primero = self._crear_concepto(documento, 'NP-1', 1)
        segundo = self._crear_concepto(documento, 'NP-2', 2)
        self._grant('change_documentoconceptos')
        self._login()

        self.client.post(
            reverse('conceptos:concepto_subir', kwargs={'pk': documento.pk, 'concepto_id': segundo.pk})
        )

        primero.refresh_from_db()
        segundo.refresh_from_db()
        self.assertEqual(primero.orden, 1)
        self.assertEqual(segundo.orden, 2)

    def test_ajax_documento_cancelado_devuelve_error(self):
        documento = self._crear_documento(status=DocumentoConceptos.STATUS_CANCELADO)
        self._crear_concepto(documento, 'NP-1', 1)
        segundo = self._crear_concepto(documento, 'NP-2', 2)
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_subir', kwargs={'pk': documento.pk, 'concepto_id': segundo.pk}),
            HTTP_X_REQUESTED_WITH='XMLHttpRequest',
        )

        self.assertEqual(response.status_code, 400)
        self.assertFalse(response.json()['ok'])
        self.assertEqual(
            response.json()['error'],
            'No se pueden reordenar conceptos de una factura cancelada.',
        )

    def test_usuario_sin_change_no_puede_reordenar(self):
        documento = self._crear_documento()
        concepto = self._crear_concepto(documento, 'NP-1', 1)
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_bajar', kwargs={'pk': documento.pk, 'concepto_id': concepto.pk})
        )

        self.assertEqual(response.status_code, 403)
        concepto.refresh_from_db()
        self.assertEqual(concepto.orden, 1)

    def test_concepto_de_otro_documento_no_puede_reordenarse(self):
        documento = self._crear_documento()
        otro_documento = self._crear_documento()
        concepto = self._crear_concepto(otro_documento, 'NP-OTRO', 1)
        self._grant('change_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:concepto_subir', kwargs={'pk': documento.pk, 'concepto_id': concepto.pk})
        )

        self.assertEqual(response.status_code, 404)

    def test_usuario_con_view_descarga_docx(self):
        documento = self._crear_documento()
        self._crear_concepto(documento, 'NP-1', 1)
        self._grant('view_documentoconceptos')
        self._login()

        response = self.client.get(
            reverse('conceptos:documento_exportar_word', kwargs={'pk': documento.pk})
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response['Content-Type'],
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        self.assertIn(f'conceptos_{documento.folio}.docx', response['Content-Disposition'])

    def test_docx_contiene_folio_conceptos_en_orden_y_total(self):
        from docx import Document

        documento = self._crear_documento()
        segundo = self._crear_concepto(documento, 'NP-2', 2)
        primero = self._crear_concepto(documento, 'NP-1', 1)
        self._grant('view_documentoconceptos')
        self._login()

        response = self.client.get(
            reverse('conceptos:documento_exportar_word', kwargs={'pk': documento.pk})
        )
        doc = Document(BytesIO(response.content))
        texto = '\n'.join(paragraph.text for paragraph in doc.paragraphs)
        tabla = doc.tables[0]
        encabezados = [cell.text for cell in tabla.rows[0].cells]
        identificacion = tabla.rows[1].cells[0].text

        self.assertIn(documento.folio, texto)
        self.assertIn('Total general: 3', texto)
        self.assertEqual(
            encabezados,
            ['Identificación', 'Descripcion', 'Cantidad', 'Precio unitario', 'Total concepto'],
        )
        self.assertNotIn('#', encabezados)
        self.assertNotIn('Numero de parte', encabezados)
        self.assertNotIn('Modelo', encabezados)
        self.assertNotIn('Serie', encabezados)
        self.assertLess(
            identificacion.index('Número de parte.'),
            identificacion.index('Modelo.'),
        )
        self.assertLess(
            identificacion.index('Modelo.'),
            identificacion.index('Serie.'),
        )
        self.assertIn(primero.numero_parte, identificacion)
        self.assertEqual(tabla.rows[1].cells[1].text, primero.descripcion)
        self.assertIn(segundo.numero_parte, tabla.rows[2].cells[0].text)

    def test_usuario_sin_view_no_puede_exportar_word(self):
        documento = self._crear_documento()
        self._login()

        response = self.client.get(
            reverse('conceptos:documento_exportar_word', kwargs={'pk': documento.pk})
        )

        self.assertEqual(response.status_code, 302)

    def test_no_quitar_concepto_de_documento_confirmado_o_cancelado(self):
        for status in (
            DocumentoConceptos.STATUS_CONFIRMADO,
            DocumentoConceptos.STATUS_CANCELADO,
        ):
            documento = self._crear_documento(status=status)
            concepto = Concepto.objects.create(
                documento=documento,
                descripcion='Sensor',
                cantidad=Decimal('1'),
                precio_unitario=Decimal('10'),
            )
            self._grant('delete_documentoconceptos')
            self._login()

            response = self.client.post(
                reverse(
                    'conceptos:concepto_delete',
                    kwargs={'pk': documento.pk, 'concepto_id': concepto.pk},
                )
            )

            self.assertEqual(response.status_code, 302)
            self.assertTrue(Concepto.objects.filter(pk=concepto.pk).exists())

    def test_confirmar_documento_requiere_permiso_custom(self):
        documento = self._crear_documento()
        self._login()

        response = self.client.post(
            reverse('conceptos:documento_confirmar', kwargs={'pk': documento.pk})
        )

        self.assertEqual(response.status_code, 302)
        documento.refresh_from_db()
        self.assertEqual(documento.status, DocumentoConceptos.STATUS_BORRADOR)

    def test_usuario_con_permiso_confirma_documento(self):
        documento = self._crear_documento()
        self._grant('puede_confirmar_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:documento_confirmar', kwargs={'pk': documento.pk})
        )

        self.assertEqual(response.status_code, 302)
        documento.refresh_from_db()
        self.assertEqual(documento.status, DocumentoConceptos.STATUS_CONFIRMADO)

    def test_confirmar_documento_crea_historial_para_conceptos_validos(self):
        documento = self._crear_documento()
        concepto = Concepto.objects.create(
            documento=documento,
            numero_parte='NP-HIS',
            serie='SER-HIS',
            modelo='MOD-HIS',
            descripcion='Sensor historial',
            cantidad=Decimal('1'),
            precio_unitario=Decimal('10'),
        )
        self._grant('puede_confirmar_documentoconceptos')
        self._login()

        self.client.post(reverse('conceptos:documento_confirmar', kwargs={'pk': documento.pk}))

        historial = HistorialCoincidencia.objects.get(concepto=concepto)
        self.assertEqual(historial.documento, documento)
        self.assertEqual(historial.confirmado_por, self.user)
        self.assertEqual(historial.regla_usada, 'manual')
        self.assertEqual(historial.match_type, 'confirmado')
        self.assertTrue(historial.usar_para_biblioteca)

    def test_confirmar_factura_despues_de_importar_genera_historial_y_patron(self):
        documento = self._crear_documento()
        NumeroParte.objects.create(
            numero_parte='NP-IMPORTADO',
            modelo='MOD-IMPORTADO',
            descripcion='Sensor importado',
            fraccion='',
            activo=True,
        )
        self._grant('change_documentoconceptos', 'puede_confirmar_documentoconceptos')
        self._login()

        self.client.post(
            reverse('conceptos:conceptos_importar', kwargs={'pk': documento.pk}),
            {
                'archivo': self._csv_upload(
                    'numero_parte,serie,modelo,descripcion,cantidad,precio_unitario\n'
                    'NP-IMPORTADO,SER-001,,,1,0\n'
                    'NP-IMPORTADO,SER-002,,,1,0\n'
                    'NP-IMPORTADO,SER-003,,,1,0\n'
                )
            },
        )
        self.client.post(
            reverse('conceptos:conceptos_importar_confirmar', kwargs={'pk': documento.pk})
        )
        self.assertFalse(HistorialCoincidencia.objects.exists())

        self.client.post(reverse('conceptos:documento_confirmar', kwargs={'pk': documento.pk}))

        self.assertEqual(HistorialCoincidencia.objects.count(), 3)
        patron = PatronSerie.objects.get(numero_parte='NP-IMPORTADO')
        self.assertEqual(patron.series_unicas, 3)
        self.assertEqual(patron.estado, PatronSerie.ESTADO_APROBADO)
        self.assertTrue(patron.activo)

    def test_confirmar_factura_actualiza_historial_no_utilizable(self):
        documento = self._crear_documento()
        concepto = Concepto.objects.create(
            documento=documento,
            numero_parte='NP-EXISTENTE',
            serie='SER-EXISTENTE',
            descripcion='Evidencia existente',
            cantidad=Decimal('1'),
            precio_unitario=Decimal('10'),
        )
        historial = HistorialCoincidencia.objects.create(
            concepto=concepto,
            documento=documento,
            numero_parte=concepto.numero_parte,
            serie=concepto.serie,
            descripcion=concepto.descripcion,
            usar_para_biblioteca=False,
        )
        self._grant('puede_confirmar_documentoconceptos')
        self._login()

        self.client.post(reverse('conceptos:documento_confirmar', kwargs={'pk': documento.pk}))

        historial.refresh_from_db()
        self.assertTrue(historial.usar_para_biblioteca)
        self.assertEqual(historial.confirmado_por, self.user)

    def test_confirmar_documento_no_duplica_historial(self):
        documento = self._crear_documento()
        concepto = Concepto.objects.create(
            documento=documento,
            numero_parte='NP-DUP',
            serie='SER-DUP',
            descripcion='Sensor duplicado',
            cantidad=Decimal('1'),
            precio_unitario=Decimal('10'),
        )
        self._grant('puede_confirmar_documentoconceptos')
        self._login()

        self.client.post(reverse('conceptos:documento_confirmar', kwargs={'pk': documento.pk}))
        self.client.post(reverse('conceptos:documento_confirmar', kwargs={'pk': documento.pk}))

        self.assertEqual(
            HistorialCoincidencia.objects.filter(concepto=concepto).count(),
            1,
        )

    def test_confirmar_documento_no_registra_conceptos_sin_serie_ni_numero_parte(self):
        documento = self._crear_documento()
        Concepto.objects.create(
            documento=documento,
            descripcion='Solo descripcion',
            cantidad=Decimal('1'),
            precio_unitario=Decimal('10'),
        )
        self._grant('puede_confirmar_documentoconceptos')
        self._login()

        self.client.post(reverse('conceptos:documento_confirmar', kwargs={'pk': documento.pk}))

        self.assertFalse(HistorialCoincidencia.objects.exists())

    def test_historial_genera_firma_normalizada(self):
        documento = self._crear_documento()
        concepto = Concepto.objects.create(
            documento=documento,
            numero_parte=' np-001 ',
            serie='ser-firma',
            modelo=' mod   a ',
            descripcion=' Sensor   de temperatura ',
            cantidad=Decimal('1'),
            precio_unitario=Decimal('10'),
        )
        self._grant('puede_confirmar_documentoconceptos')
        self._login()

        self.client.post(reverse('conceptos:documento_confirmar', kwargs={'pk': documento.pk}))

        historial = HistorialCoincidencia.objects.get(concepto=concepto)
        self.assertEqual(
            historial.firma_json,
            {
                'numero_parte': 'NP-001',
                'modelo': 'MOD A',
                'descripcion': 'SENSOR DE TEMPERATURA',
            },
        )
        self.assertEqual(
            historial.firma_texto,
            'numero_parte=NP-001|modelo=MOD A|descripcion=SENSOR DE TEMPERATURA',
        )

    def test_usuario_con_permiso_cancela_documento(self):
        documento = self._crear_documento()
        self._grant('puede_cancelar_documentoconceptos')
        self._login()

        response = self.client.post(
            reverse('conceptos:documento_cancelar', kwargs={'pk': documento.pk})
        )

        self.assertEqual(response.status_code, 302)
        documento.refresh_from_db()
        self.assertEqual(documento.status, DocumentoConceptos.STATUS_CANCELADO)
